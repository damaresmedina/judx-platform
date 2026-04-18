"""
md-to-docx.py — converte markdown para docx usando python-docx.

Suporta:
- headings h1-h6
- parágrafos
- tabelas (| ... |)
- listas bullet (- ou *)
- listas numeradas (1. 2. ...)
- code blocks (```)
- código inline (`x`)
- bold (**x**) e italic (*x* ou _x_)

Uso: python md-to-docx.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_inline(p, text):
    """Parse inline (bold, italic, code) e adiciona runs ao paragraph."""
    # Tokens: **bold**, `code`, *italic* ou _italic_
    pattern = re.compile(r'(\*\*[^\*]+\*\*|`[^`]+`|\*[^\*]+\*|_[^_]+_)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            p.add_run(text[last:m.start()])
        token = m.group(1)
        if token.startswith('**') and token.endswith('**'):
            r = p.add_run(token[2:-2]); r.bold = True
        elif token.startswith('`') and token.endswith('`'):
            r = p.add_run(token[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
            r = p.add_run(token[1:-1]); r.italic = True
        else:
            p.add_run(token)
        last = m.end()
    if last < len(text):
        p.add_run(text[last:])

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    # Margens razoáveis
    for sec in doc.sections:
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    i = 0
    while i < len(lines):
        line = lines[i]
        # ---- Code blocks ----
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Consolas'; r.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        # ---- Headings ----
        h = re.match(r'^(#{1,6})\s+(.+)$', line)
        if h:
            level = len(h.group(1))
            text = h.group(2).strip()
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue
        # ---- Tables ----
        if line.strip().startswith('|') and i+1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
            # cabeçalho | sep | linhas
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells); i += 1
            t = doc.add_table(rows=1+len(rows), cols=len(header))
            t.style = 'Light Grid Accent 1'
            for ci, h in enumerate(header):
                cell = t.rows[0].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]; parse_inline(p, h)
                for run in p.runs: run.bold = True
            for ri, row in enumerate(rows):
                for ci, cv in enumerate(row):
                    if ci < len(t.rows[ri+1].cells):
                        cell = t.rows[ri+1].cells[ci]
                        cell.text = ''
                        parse_inline(cell.paragraphs[0], cv)
            doc.add_paragraph()  # espaçamento
            continue
        # ---- Horizontal rule ----
        if re.match(r'^\s*---+\s*$', line):
            doc.add_paragraph('').add_run('')
            i += 1
            continue
        # ---- Bullet list ----
        m = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, m.group(2))
            i += 1
            continue
        # ---- Numbered list ----
        m = re.match(r'^\s*\d+\.\s+(.+)$', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, m.group(1))
            i += 1
            continue
        # ---- Parágrafo ou vazio ----
        if line.strip():
            # quebra linha seguida vira parágrafo separado
            p = doc.add_paragraph()
            parse_inline(p, line)
        else:
            # linha em branco = espaçamento
            pass
        i += 1

    doc.save(docx_path)
    print(f'saved: {docx_path}')

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('uso: python md-to-docx.py <in.md> <out.docx>'); sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
