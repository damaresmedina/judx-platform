"""
Gera relatório Word com todos os achados do banco JudX/ICONS — 27/03/2026
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ───────────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x2a, 0x44)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

# ── Cover ────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RELATÓRIO DE ACHADOS EMPÍRICOS')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Bancos JudX e ICONS — Análise Completa')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('27 de março de 2026')
run.font.size = Pt(14)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJUS · Instituto Constituição Aberta — ICONS')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Corpus: 169.851 decisões STF · 1.420 temas repetitivos STJ · 2.509 processos-semente')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── TOC ──────────────────────────────────────────────────

doc.add_heading('Sumário', level=1)
toc_items = [
    '1. Infraestrutura de Dados',
    '2. Ancoragem Constitucional — ICONS',
    '3. STJ — Temas Repetitivos (Camada 1A)',
    '4. Mapa de Contaminação STJ',
    '5. Taxa de Não-Decisão',
    '6. Análise de Ambiente — STF Virtual',
    '7. Coalizões e Divergência',
    '8. Contrafactual Marco Aurélio',
    '9. Impacto Econômico — Dados Externos',
    '10. Síntese dos Achados',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. INFRAESTRUTURA
# ══════════════════════════════════════════════════════════

doc.add_heading('1. Infraestrutura de Dados', level=1)

doc.add_paragraph(
    'O sistema opera sobre dois bancos de dados independentes em Supabase (PostgreSQL): '
    'o banco JudX (comportamento institucional de STF e STJ) e o banco ICONS (ancoragem constitucional do STF). '
    'Os dados foram extraídos de fontes públicas oficiais: Corte Aberta do STF, portal de repetitivos do STJ, '
    'CKAN do STJ, e API Pública Datajud do CNJ.'
)

add_table(
    ['Base', 'Registros', 'Fonte', 'Status'],
    [
        ['stf_decisoes (raw)', '169.851', 'Corte Aberta STF', 'Completo'],
        ['judx_case (normalizado)', '139.737', 'Pipeline JudX', 'Completo'],
        ['judx_decision', '224.887', 'Pipeline JudX', 'Completo'],
        ['stf_partes', '~140.000', 'Portal STF', 'Em extração (15%)'],
        ['stj_temas', '1.420', 'Portal Repetitivos STJ', 'Completo'],
        ['stj_processos_semente', '2.509', 'Portal Repetitivos STJ', 'Completo'],
        ['stj_contramostra', '3.902', 'CKAN STJ + Datajud CNJ', 'Completo'],
        ['ICONS edges', '7.766', 'Ancoragem normativa', 'Completo'],
        ['ICONS objects', '126.545', 'Registros jurisprudenciais', 'Completo'],
    ]
)

add_note('Nota: 169.851 decisões raw produzem 139.737 cases únicos porque múltiplas decisões incidem sobre o mesmo processo (ON CONFLICT).')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. ANCORAGEM ICONS
# ══════════════════════════════════════════════════════════

doc.add_heading('2. Ancoragem Constitucional — ICONS', level=1)

doc.add_paragraph(
    'O banco ICONS contém 7.766 edges do tipo ancora_normativa, conectando registros jurisprudenciais '
    'a dispositivos da Constituição Federal de 1988. A validação confirmou que 100% dos source_id são '
    'registros jurisprudenciais reais — nenhum dado espúrio. A granularidade alcança o nível de alínea.'
)

doc.add_heading('Top 10 dispositivos por densidade decisória', level=2)

add_table(
    ['Dispositivo', 'Decisões vinculadas', 'Matéria'],
    [
        ['CF/88 Art. 5º', '1.049', 'Direitos fundamentais'],
        ['CF/88 Art. 37', '409', 'Administração pública'],
        ['CF/88 Art. 93', '86', 'Estatuto da magistratura'],
        ['CF/88 Art. 40', '79', 'Previdência do servidor'],
        ['CF/88 Art. 2º', '76', 'Separação de poderes'],
        ['CF/88 Art. 71', '61', 'Tribunal de Contas'],
        ['CF/88 Art. 102', '58', 'Competência do STF'],
        ['CF/88 Art. 100', '54', 'Precatórios'],
        ['CF/88 Art. 196', '48', 'Direito à saúde'],
        ['CF/88 Art. 195', '44', 'Seguridade social'],
    ]
)

add_note('Art. 5º concentra 13,5% de todas as ancoragens normativas — domínio absoluto dos direitos fundamentais no contencioso constitucional.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. STJ TEMAS
# ══════════════════════════════════════════════════════════

doc.add_heading('3. STJ — Temas Repetitivos (Camada 1A)', level=1)

doc.add_paragraph(
    '1.420 temas repetitivos extraídos com parser de 15 páginas do portal de repetitivos do STJ. '
    'Zero erros de extração. Cada tema inclui: situação, órgão julgador, ramo do direito, questão submetida, '
    'tese firmada, processos afetados, relator, datas e link para tema STF quando há repercussão geral.'
)

doc.add_heading('Distribuição por situação', level=2)

add_table(
    ['Situação', 'Quantidade', '%'],
    [
        ['Trânsito em julgado', '973', '68,5%'],
        ['Cancelado', '190', '13,4%'],
        ['Afetado', '113', '8,0%'],
        ['Julgado (sem trânsito)', '90', '6,3%'],
        ['Em Julgamento', '22', '1,5%'],
        ['Revisado', '15', '1,1%'],
        ['Sobrestado', '14', '1,0%'],
    ]
)

doc.add_heading('Top ramos do direito', level=2)

add_table(
    ['Ramo', 'Temas', 'Dias médio afetação→julgamento'],
    [
        ['Processual Civil e Trabalho', '420', '419'],
        ['Tributário', '292', '329'],
        ['Administrativo', '248', '254'],
        ['Civil', '174', '386'],
        ['Previdenciário', '88', '421'],
        ['Penal', '77', '272'],
        ['Processual Penal', '60', '318'],
        ['Consumidor', '55', '389'],
    ]
)

doc.add_paragraph(
    'Dados-chave: 1.094 temas (77%) têm tese firmada. 257 temas possuem link para tema de Repercussão Geral '
    'no STF — estes são os circuitos de circularidade STJ↔STF documentados na Parte II do paper.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. MAPA DE CONTAMINAÇÃO
# ══════════════════════════════════════════════════════════

doc.add_heading('4. Mapa de Contaminação STJ', level=1)

doc.add_heading('4.1 Vetores humanos — relatores que mais afetam', level=2)

add_table(
    ['Relator', 'Temas', 'Trânsito', 'Pendentes', 'Dias médio'],
    [
        ['LUIZ FUX', '131', '121', '0', '132'],
        ['MAURO CAMPBELL MARQUES', '108', '87', '1', '503'],
        ['LUIS FELIPE SALOMÃO', '97', '87', '0', '382'],
        ['NAPOLEÃO NUNES MAIA FILHO', '75', '51', '0', '584'],
        ['HERMAN BENJAMIN', '68', '54', '1', '448'],
        ['OG FERNANDES', '59', '35', '0', '631'],
        ['MARIA THEREZA DE ASSIS MOURA', '57', '27', '5', '195'],
        ['MARCO AURÉLIO BELLIZZE', '32', '10', '9', '539'],
    ]
)

add_note('Luiz Fux afeta mais e resolve mais rápido (132 dias). Bellizze tem o maior backlog ativo (9 pendentes, 539 dias médio).')

doc.add_heading('4.2 Tribunais de origem — onde nasce o tema repetitivo', level=2)

add_table(
    ['Tribunal', 'Processos-semente', '% do total'],
    [
        ['TRF4 (RS/SC/PR)', '388', '15,5%'],
        ['TJSP (combinado)', '311', '12,4%'],
        ['TRF3 (SP/MS)', '251', '10,0%'],
        ['TRF5 (NE)', '247', '9,8%'],
        ['TJRS', '214', '8,5%'],
        ['TRF1 (DF/GO/MG)', '173', '6,9%'],
        ['TJMG', '143', '5,7%'],
        ['TJPR', '114', '4,5%'],
        ['TJRJ', '96', '3,8%'],
    ]
)

add_note('O TRF4 é o maior gerador de temas repetitivos do STJ — o circuito Sul alimenta desproporcionalmente o tribunal superior.')

doc.add_heading('4.3 Concentração geográfica por ramo', level=2)

doc.add_paragraph(
    'SP domina em Processual Civil (188 sementes) e Tributário (150). '
    'RS é desproporcionalmente presente em Administrativo (95) e Civil (87). '
    'MG concentra Processual Penal (35). A geografia da contaminação não é uniforme — cada UF alimenta o STJ em ramos específicos.'
)

doc.add_heading('4.4 Órgão julgador × Ramo', level=2)

add_table(
    ['Órgão', 'Ramo', 'Temas', 'Dias médio'],
    [
        ['Primeira Seção', 'Tributário', '290', '329'],
        ['Primeira Seção', 'Administrativo', '218', '260'],
        ['Primeira Seção', 'Processual Civil', '202', '407'],
        ['Segunda Seção', 'Civil', '151', '368'],
        ['Corte Especial', 'Processual Civil', '106', '480'],
        ['Terceira Seção', 'Penal', '77', '272'],
    ]
)

add_note('A Primeira Seção concentra 71% dos temas. A Corte Especial é a mais lenta (480 dias médio).')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. TAXA DE NÃO-DECISÃO
# ══════════════════════════════════════════════════════════

doc.add_heading('5. Taxa de Não-Decisão', level=1)

doc.add_heading('5.1 STF — 79% sem mérito', level=2)

doc.add_paragraph(
    'A análise das 169.851 decisões do corpus Corte Aberta do STF revela que mais de 79% '
    'das decisões não apreciam o mérito da causa. A maioria termina como "Agravo regimental não provido", '
    '"Embargos rejeitados" ou "Negado seguimento" — mecanismos de filtragem processual que não produzem '
    'vinculação substantiva.'
)

doc.add_heading('5.2 STJ — 90% sem mérito', level=2)

doc.add_paragraph(
    'Análise de um dia de metadados do CKAN STJ (19/03/2026, 5.968 decisões):'
)

add_table(
    ['Classe', '% do fluxo', 'Negando', 'Não conhecido', '% sem mérito'],
    [
        ['AREsp', '74,7%', '58,3%', '36,7%', '95,0%'],
        ['REsp', '14,7%', '50,9%', '20,0%', '70,9%'],
        ['HC', '7,0%', '47,4%', '38,5%', '85,9%'],
    ]
)

doc.add_paragraph(
    'O AREsp (Agravo em Recurso Especial) é 75% do volume diário do STJ e termina com 95% de taxa '
    'de não-decisão de mérito. Ponderando todas as classes, aproximadamente 90% do STJ termina sem apreciação de mérito — '
    'taxa superior aos 79% do STF.'
)

add_note('Fonte: CKAN STJ, dataset "Íntegras de Decisões Terminativas e Acórdãos do Diário da Justiça", metadados de 19/03/2026. Disponível a partir de fev/2022.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. AMBIENTE VIRTUAL
# ══════════════════════════════════════════════════════════

doc.add_heading('6. Análise de Ambiente — STF Virtual', level=1)

doc.add_heading('6.1 Descoberta estrutural', level=2)

doc.add_paragraph(
    'Todas as decisões colegiadas do STF no período 2016-2025 são registradas como virtuais '
    '(campo decisoes_virtual = true). O campo "Presencial" marca apenas decisões monocráticas. '
    'O STF opera em sessão virtual como regime padrão para decisões colegiadas.'
)

doc.add_heading('6.2 Volume por sessão — prova da assessorização', level=2)

add_table(
    ['Semana', 'Processos', 'Relatores ativos', 'Média por relator'],
    [
        ['25/ago/2025', '1.377', '13', '106'],
        ['19/ago/2024', '1.347', '11', '123'],
        ['17/fev/2025', '1.339', '11', '122'],
        ['21/ago/2023', '1.325', '12', '110'],
        ['24/ago/2020', '1.187', '12', '99'],
    ]
)

doc.add_paragraph(
    '122 processos por relator por semana — 24 acórdãos colegiados por dia útil. '
    'O número é incompatível com deliberação individual real e evidencia a produção '
    'por assessoria (assessorização da jurisdição).'
)

doc.add_heading('6.3 Unanimidade por órgão', level=2)

add_table(
    ['Órgão', 'Total acórdãos', '% Unânime', 'Por maioria'],
    [
        ['2ª Turma', '46.487', '91,3%', '4.064'],
        ['Tribunal Pleno', '44.986', '85,5%', '7.164'],
        ['1ª Turma', '51.796', '84,1%', '10.166'],
        ['Plenário Virtual RG', '1.382', '41,0%', '513'],
    ]
)

add_note('O Plenário Virtual de Repercussão Geral é o único espaço com divergência real frequente (59% não unânime). Turmas e Pleno operam em regime de pseudo-colegialidade.')

doc.add_heading('6.4 Tamanho do texto — proxy de profundidade deliberativa', level=2)

add_table(
    ['Tipo de resultado', 'Tam. médio (caracteres)', 'Natureza'],
    [
        ['Procedente (mérito)', '1.162', 'Decisão real'],
        ['Recebida denúncia', '557', 'Decisão penal'],
        ['Agravo provido', '479', 'Reversão'],
        ['Embargos recebidos', '400', 'Correção'],
        ['Agravo não provido', '~190 (mediana geral)', 'Filtro processual'],
    ]
)

doc.add_paragraph(
    'A mediana do texto ficou estável em ~190 caracteres de 2016 a 2025. '
    'Os acórdãos virtuais não ficaram mais curtos — sempre foram curtos. '
    'A maioria se resume a "nos termos do voto do Relator" sem fundamentação adicional.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. COALIZÕES
# ══════════════════════════════════════════════════════════

doc.add_heading('7. Coalizões e Divergência', level=1)

doc.add_heading('7.1 Série histórica da divergência', level=2)

add_table(
    ['Ano', 'Acórdãos', '% divergência', 'Nota'],
    [
        ['2016', '5.062', '24,4%', ''],
        ['2017', '10.108', '16,3%', ''],
        ['2018', '11.665', '16,9%', ''],
        ['2019', '14.458', '13,2%', ''],
        ['2020', '17.192', '22,8%', 'Pandemia + matérias explosivas'],
        ['2021', '15.023', '20,3%', 'Marco Aurélio aposenta jul/2021'],
        ['2022', '12.706', '6,1%', 'ANOMALIA — menor da série'],
        ['2023', '17.910', '14,0%', 'Inquéritos 8 de janeiro'],
        ['2024', '20.921', '10,7%', ''],
        ['2025', '18.773', '11,6%', ''],
    ]
)

doc.add_heading('7.2 Ministros mais vencidos', level=2)

add_table(
    ['Ministro', 'Vezes vencido', 'Nota'],
    [
        ['Marco Aurélio', '8.451', 'Maior dissidente da história do STF'],
        ['Mendonça + Nunes (juntos)', '2.520', 'Bloco de minoria estável'],
        ['Edson Fachin', '1.146', ''],
        ['Gilmar Mendes', '840', ''],
        ['Flávio Dino', '606', 'Em apenas 2 anos'],
        ['Alexandre de Moraes', '411', ''],
        ['Fachin (como Relator)', '280', 'Relator de fato derrotado'],
    ]
)

doc.add_heading('7.3 Bloco Mendonça + Nunes Marques', level=2)

doc.add_paragraph(
    'O bloco Mendonça + Nunes Marques é uma minoria estável que vota junto em 67% das vezes. '
    'Sua derrota concentra-se em matéria penal/processual penal (62% das vezes em que são vencidos). '
    'Alexandre de Moraes é o relator em 54% dos casos onde este bloco perde.'
)

add_table(
    ['Ano', 'Vezes vencidos'],
    [
        ['2021', '128'],
        ['2022', '220'],
        ['2023', '1.833'],
        ['2024', '781'],
        ['2025', '715'],
    ]
)

add_note('Explosão em 2023 (8× o volume de 2022) — coincide com inquéritos de 8 de janeiro e operações contra milícias digitais.')

doc.add_heading('7.4 Ramos com mais divergência', level=2)

add_table(
    ['Ramo', 'Total acórdãos', '% divergência'],
    [
        ['Trabalho', '8.627', '28,9%'],
        ['Alta Complexidade/RG', '1.203', '26,2%'],
        ['Processual Penal', '32.277', '18,8%'],
        ['Previdenciário', '3.677', '18,2%'],
        ['Processual Civil', '12.379', '15,0%'],
        ['Administrativo', '42.028', '12,2%'],
        ['Tributário', '14.955', '11,4%'],
        ['Civil', '7.179', '9,0%'],
    ]
)

add_note('Trabalho é o ramo mais litigioso do STF — quase 1 em 3 acórdãos tem divergência. Civil é o mais consensual.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. CONTRAFACTUAL MARCO AURÉLIO
# ══════════════════════════════════════════════════════════

doc.add_heading('8. Contrafactual Marco Aurélio', level=1)

doc.add_paragraph(
    'Para testar se a anomalia de 2020-2021 (divergência de 22,8% e 20,3%) é explicável por um único ministro, '
    'removemos Marco Aurélio da série e recalculamos a taxa de divergência.'
)

add_table(
    ['Ano', 'Total', '% com MA', '% sem MA', 'Delta (pp)'],
    [
        ['2018', '11.665', '16,9%', '16,8%', '0,1'],
        ['2019', '14.458', '13,2%', '13,1%', '0,1'],
        ['2020', '17.192', '22,8%', '21,7%', '1,1'],
        ['2021', '15.023', '20,3%', '19,2%', '1,1'],
        ['2022', '12.706', '6,1%', '6,0%', '0,1'],
        ['2023', '17.910', '14,0%', '13,9%', '0,1'],
        ['2024', '20.921', '10,7%', '10,7%', '0,0'],
        ['2025', '18.773', '11,6%', '11,6%', '0,0'],
    ]
)

doc.add_paragraph(
    'Resultado: Marco Aurélio NÃO explica a anomalia. O delta é de apenas 1,1 ponto percentual em 2020-2021. '
    'Removendo-o completamente, a divergência cai de 22,8% para 21,7% — diferença residual. '
    'A anomalia 2020-2021 é genuína e sistêmica — provavelmente associada ao impacto da pandemia '
    'sobre matérias constitucionais (saúde, liberdades, federalismo) e à virtualização compulsória. '
    'A queda para 6,0% em 2022 permanece a verdadeira anomalia da série, e não é explicável '
    'pela composição do tribunal.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. IMPACTO ECONÔMICO
# ══════════════════════════════════════════════════════════

doc.add_heading('9. Impacto Econômico — Dados Externos', level=1)

doc.add_heading('9.1 Pesquisa FIESP — Obstáculos ao Crescimento (2023)', level=2)

doc.add_paragraph(
    '413 empresas industriais de SP consultadas (abr-mai/2023). '
    'Tributação é o obstáculo #1, Segurança Jurídica é o #3 entre 14 grandes temas avaliados.'
)

add_table(
    ['Obstáculo tributário', '% menções'],
    [
        ['Complexidade da legislação do ICMS e frequentes alterações', '23,9%'],
        ['Descasamento pagamento ICMS vs recebimento vendas', '20,1%'],
        ['Altos custos do regime ICMS-ST', '14,1%'],
        ['Gastos para preparar e pagar ICMS', '8,6%'],
    ]
)

add_table(
    ['Obstáculo judicial', 'Total', 'Empresas >R$30mi'],
    [
        ['Excesso de normas e regulamentos', '37,3%', '36,5%'],
        ['Morosidade do sistema judiciário de SP', '15,7%', '20,1%'],
        ['Custo e acessibilidade do judiciário', '8,3%', '8,4%'],
    ]
)

add_note('Para empresas acima de R$30mi, morosidade judicial sobe para 20,1% — quanto maior a empresa, mais o Judiciário pesa como obstáculo.')

doc.add_heading('9.2 Teto Decorativo — Honorários de Sucumbência (2025)', level=2)

doc.add_paragraph(
    'Estudo da Transparência Brasil + Movimento Pessoas à Frente (dez/2025) sobre o Conselho Curador '
    'dos Honorários Advocatícios (CCHA):'
)

add_table(
    ['Métrica', 'Valor'],
    [
        ['Total pago pelo CCHA (2020-ago/2025)', 'R$ 12,7 bilhões'],
        ['Beneficiários', '13.242 advogados/procuradores'],
        ['Pagamentos extrateto em 2025 (8 meses)', 'R$ 3,8 bilhões'],
        ['% ativos acima do teto em 2025', '93%'],
        ['Mediana por pessoa em jul/2025', 'R$ 304 mil'],
        ['Beneficiários que receberam >R$1 milhão', '7.649 (58%)'],
    ]
)

doc.add_paragraph(
    'O STF (ADI 6.053) e o TCU (Acórdão 307/21) determinam que honorários de sucumbência têm natureza '
    'remuneratória e devem respeitar o teto constitucional. O CCHA contorna essa determinação classificando '
    'pagamentos como indenizatórios — criando "penduricalhos" como auxílio-saúde complementar, '
    'complementação de férias retroativa a 2017, e rateio extraordinário.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. SÍNTESE
# ══════════════════════════════════════════════════════════

doc.add_heading('10. Síntese dos Achados', level=1)

findings = [
    ('O sistema não decide',
     '79% do STF e ~90% do STJ terminam sem apreciação de mérito. O produto majoritário '
     'do sistema judicial superior brasileiro é a não-decisão.'),
    ('O virtual não delibera',
     '85-91% de unanimidade nas sessões virtuais, com 122 processos por relator por semana nas sessões '
     'mais pesadas. O volume é incompatível com deliberação individual.'),
    ('A divergência é concentrada',
     'Trabalho (28,9% de divergência) vs. Civil (9,0%). Processual Penal (18,8%) é o campo de batalha '
     'do bloco Mendonça + Nunes Marques.'),
    ('Há uma minoria estável',
     'O bloco Mendonça + Nunes Marques vota junto em 67% das vezes, é vencido primariamente em matéria penal, '
     'e Alexandre de Moraes relata 54% dos casos onde esse bloco perde.'),
    ('2022 é uma anomalia',
     '6,1% de divergência — menos da metade de qualquer outro ano da série. Não é explicável pela '
     'composição do tribunal (contrafactual Marco Aurélio demonstra delta residual de 0,1pp).'),
    ('O custo da incerteza é mensurável',
     '292 temas tributários repetitivos × 329 dias médios de resolução no STJ. A FIESP documenta '
     'complexidade tributária como obstáculo #1 ao crescimento industrial.'),
    ('O TRF4 é o maior gerador de contaminação',
     '388 processos-semente (15,5% do total) originam-se do TRF4. O circuito Sul alimenta desproporcionalmente '
     'o sistema de temas repetitivos do STJ.'),
    ('Os honorários crescem com a litigiosidade',
     'R$ 3,8 bilhões pagos extrateto em 2025, financiados pela sucumbência que o sistema produz ao não decidir. '
     'O circuito de incentivos é auto-referente: quanto mais o sistema não decide, mais recursos de sucumbência circulam.'),
]

for i, (titulo, texto) in enumerate(findings, 1):
    doc.add_heading(f'Achado {i}: {titulo}', level=2)
    doc.add_paragraph(texto)

doc.add_page_break()

# ── Methodology ──────────────────────────────────────────

doc.add_heading('Nota Metodológica', level=1)

doc.add_paragraph(
    'Os dados apresentados derivam de dois corpora distintos e complementares:'
)

items = [
    'Corpus JudX/STF: 169.851 decisões extraídas do Corte Aberta do STF, normalizadas em 139.737 cases e 224.887 decisions.',
    'Corpus STJ: 1.420 temas repetitivos extraídos do portal de repetitivos, com 2.509 processos-semente e 3.902 processos de contramostra.',
    'Banco ICONS: 7.766 edges de ancoragem normativa conectando registros jurisprudenciais a dispositivos da CF/88.',
    'CKAN STJ: metadados diários de decisões terminativas (fev/2022 em diante).',
    'API Pública Datajud (CNJ): processos do STJ de 2008 a 2021.',
    'Dados externos: Pesquisa FIESP/DECOMTEC (jun/2023), Estudo Transparência Brasil (dez/2025).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'A "divergência" é medida pela presença da string "vencido" no campo observacao_andamento das decisões colegiadas. '
    'Esta é uma medida conservadora que captura qualquer ministro vencido (não apenas o relator). '
    'Para identificação precisa do relator vencido, o padrão "Relator" entre parênteses após o nome '
    'do ministro é utilizado como filtro adicional.'
)

doc.add_paragraph(
    'O campo "ambiente de julgamento" no corpus Corte Aberta marca decisoes_virtual = true para todas as '
    'decisões colegiadas (2016-2025). O campo ambiente_julgamento (Presencial/Virtual) só está presente '
    'para decisões a partir de 2026. A distinção virtual/presencial nas análises refere-se ao flag '
    'decisoes_virtual e deve ser interpretada com esta limitação.'
)

# ── Save ─────────────────────────────────────────────────

output_path = r'C:\Users\medin\Desktop\Relatorio_Achados_JudX_ICONS_27mar2026.docx'
doc.save(output_path)
print(f'Relatório salvo em: {output_path}')
print(f'Páginas estimadas: ~25')
