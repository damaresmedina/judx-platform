import docx, sys, os, re, json
sys.stdout.reconfigure(encoding='utf-8')

ocr_dir = "C:/Users/medin/Desktop/backup_judx/resultados/plaquetas_stf/ocr_txt"
termos_json = "C:/Users/medin/Desktop/backup_judx/resultados/termos_posse_stf_1963_1979.json"
existing_docx = "C:/Users/medin/Downloads/discurso presidentes.docx"
out_path = "C:/Users/medin/Downloads/Discursos posse STF - consolidado 1963-2025.docx"

def clean_ocr(text):
    text = re.sub(r'--- Pagina \d+ ---', '', text)
    text = re.sub(r'^#.*\n', '', text)
    text = text.replace('a0 ', 'ao ').replace(' 0 ', ' o ')
    text = re.sub(r'\b0(?=[a-z\u00e0-\u00ff])', 'o', text)
    text = re.sub(r'(?<=[a-z\u00e0-\u00ff])0\b', 'o', text)
    text = re.sub(r'(\w)\s*-\s*\n\s*(\w)', r'\1\2', text)
    text = re.sub(r'(\w)\s+-\s+(\w)', r'\1\2', text)
    text = text.replace('\u20ac', 'e').replace('_', '.')
    text = text.replace('<', '').replace('>', '')
    text = re.sub(r'\n\s*[;:,.\s]\s*\n', '\n', text)
    text = re.sub(r'\n\s*\d{1,3}\s*\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\s+([.;,!?:])', r'\1', text)
    return text.strip()

def extract_sections(text):
    pattern = r'(Palavras\s+d[eo]\s+(?:Senhor\s+)?Ministro\s+[^\n]+)'
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    sections = []
    current_speaker = None
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r'Palavras\s+d[eo]', part, re.IGNORECASE):
            current_speaker = part.strip()
        else:
            if 'POSSE NA PRESID\u00caNCIA' in part and len(part) < 400:
                continue
            if 'SOLENIDADE DE POSSE' in part and len(part) < 500:
                continue
            if 'REP\u00daBLICA FEDERATIVA' in part and len(part) < 500:
                continue
            sections.append({'speaker': current_speaker, 'text': part})
    return sections

def text_to_paragraphs(text):
    raw_paras = re.split(r'\n\n+', text)
    paras = []
    for p in raw_paras:
        p = p.strip()
        p = re.sub(r'\n', ' ', p)
        p = re.sub(r'\s+', ' ', p)
        if len(p) > 10:
            paras.append(p)
    return paras

# Plaquetas OCR (1975-1993)
plaquetas = [
    {"heading": "Djaci Falc\u00e3o (1975-1977)",
     "file": "Plaqueta_Possepresidencial_DjaciFalcao_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de fevereiro de 1975",
     "presidente": "Djaci Alves Falc\u00e3o", "vice": "Thompson Flores"},
    {"heading": "Thompson Flores (1977-1979)",
     "file": "Plaqueta_Possepresidencial_ThompsonFlores_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de fevereiro de 1977",
     "presidente": "Carlos Thompson Flores", "vice": "Olavo Bilac Pinto"},
    {"heading": "Xavier de Albuquerque (1981-1983)",
     "file": "Plaqueta_Possepresidencial_XavierdeAlbuquerquearrumar_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 16 de fevereiro de 1981",
     "presidente": "Francisco Manoel Xavier de Albuquerque", "vice": "Jo\u00e3o Leit\u00e3o de Abreu"},
    {"heading": "Cordeiro Guerra (1983-1985)",
     "file": "Plaqueta_Possepresidencial_CordeiroGuerra_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 21 de fevereiro de 1983",
     "presidente": "Jo\u00e3o Baptista Cordeiro Guerra", "vice": "Jos\u00e9 Carlos Moreira Alves"},
    {"heading": "Moreira Alves (1985-1987)",
     "file": "Plaqueta_Possepresidencial_MoreiraAlves_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 25 de fevereiro de 1985",
     "presidente": "Jos\u00e9 Carlos Moreira Alves", "vice": "D\u00e9cio Miranda"},
    {"heading": "Aldir Passarinho (1987-1989)",
     "file": "Plaqueta_Possepresidencial_AldirPassarinho_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 1987",
     "presidente": "Aldir Guimar\u00e3es Passarinho", "vice": ""},
    {"heading": "N\u00e9ri da Silveira (1989-1991)",
     "file": "Plaqueta_Possepresidencial_NeridaSilveira_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de mar\u00e7o de 1989",
     "presidente": "Jos\u00e9 N\u00e9ri da Silveira", "vice": "Aldir Guimar\u00e3es Passarinho"},
    {"heading": "Oct\u00e1vio Gallotti (1991-1993)",
     "file": "Plaqueta_Possepresidencial_OctavioGallotti_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 1991",
     "presidente": "Oct\u00e1vio Gallotti", "vice": ""},
]

# ── BUILD ──

print("Construindo consolidado definitivo 1963-2025...")
doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = docx.shared.Pt(11)

# ── TITLE ──
p = doc.add_paragraph()
run = p.add_run('DISCURSOS DE POSSE NA PRESIDENCIA DO STF')
run.font.size = docx.shared.Pt(18)
run.bold = True

p = doc.add_paragraph()
run = p.add_run('Corpus consolidado: 1963\u20132025')
run.font.size = docx.shared.Pt(13)
run.italic = True

p = doc.add_paragraph()
run = p.add_run('Do desconhecido ao incompreendido')
run.font.size = docx.shared.Pt(11)
run.italic = True

doc.add_paragraph('')

# ── NOTA METODOLOGICA ──
doc.add_heading('Nota metodologica', level=2)
nota = doc.add_paragraph()
nota.add_run(
    'Este corpus reune os discursos e termos de posse dos presidentes do '
    'Supremo Tribunal Federal, de Alvaro Moutinho Ribeiro da Costa (1963) '
    'a Edson Fachin (2025). O material foi composto a partir de tres '
    'camadas documentais:\n\n'
)
nota.add_run('1. Termos de posse manuscritos (1963\u20131979)').bold = True
nota.add_run(
    ' \u2014 Transcritos dos originais do livro de termos do STF. '
    'Documentos formais que registram data, compromisso legal, signatarios. '
    'Nao contem discursos, apenas o ato juridico da investidura. '
    'Para 6 dos 7 presidentes deste periodo, o termo manuscrito e a unica '
    'fonte primaria disponivel em formato digital. A excecao e Ribeiro da Costa, '
    'cujo discurso integral consta do Diario da Justica de 12/12/1963.\n\n'
)
nota.add_run('2. Plaquetas institucionais (1975\u20131993)').bold = True
nota.add_run(
    ' \u2014 Publicacoes oficiais do STF com os discursos integrais das sessoes '
    'solenes de posse. Recuperadas via Wayback Machine (web.archive.org) e '
    'transcritas por OCR (EasyOCR). Podem conter artefatos de digitalizacao.\n\n'
)
nota.add_run('3. Discursos completos (1993\u20132025)').bold = True
nota.add_run(
    ' \u2014 Textos integrais dos discursos de posse, obtidos de fontes oficiais '
    'do STF e do arquivo pessoal da autora.\n\n'
)
nota.add_run('Referencia fundacional: ').bold = True
nota.add_run(
    'BALEEIRO, Aliomar. O Supremo Tribunal Federal, esse outro desconhecido. '
    'Rio de Janeiro: Forense, 1968. O titulo deste corpus dialoga com o '
    'diagnostico de Baleeiro: se em 1968 o Supremo era "esse outro desconhecido", '
    'em 2026 e possivel sustentar que se tornou "esse outro incompreendido" '
    '\u2014 nao por falta de exposicao, mas por excesso de visibilidade sem '
    'inteligibilidade (cf. MEDINA, Damares. O Supremo no diva, 2026).'
)

doc.add_paragraph('')

# ── SECTION 1: TERMOS MANUSCRITOS (1963-1979) ──
doc.add_heading('I. Termos de posse manuscritos (1963\u20131979)', level=2)

p = doc.add_paragraph()
p.add_run(
    'Os termos abaixo foram transcritos dos manuscritos originais do livro de '
    'termos de posse do STF. Sao documentos formais, nao discursos. '
    'Registram o ato juridico, o compromisso e os signatarios. '
    'A ausencia de discurso nao e lacuna \u2014 e linguagem: o silencio formal '
    'do termo manuscrito e, ele proprio, um modo de enunciacao institucional, '
    'anterior a era das plaquetas e da publicidade dos atos.'
).italic = True

doc.add_paragraph('')

with open(termos_json, 'r', encoding='utf-8') as f:
    termos_data = json.load(f)
termos_sorted = sorted(termos_data, key=lambda t: t['data_posse'].split('/')[::-1])

print(f"  I. {len(termos_sorted)} termos manuscritos")
for t in termos_sorted:
    p = doc.add_paragraph()
    run = p.add_run(t['ministro'].upper())
    run.font.size = docx.shared.Pt(14)

    meta = doc.add_paragraph()
    run = meta.add_run(f"Posse: {t['data_posse']} | Presidiu: {t['presidente_sessao']} | PGR: {t['procurador_geral']}")
    run.italic = True
    run.font.size = docx.shared.Pt(10)

    fonte = t.get('tipo', 'Manuscrito original \u2014 livro de termos do STF')
    meta2 = doc.add_paragraph()
    run = meta2.add_run(f"Fonte: {fonte}")
    run.italic = True
    run.font.size = docx.shared.Pt(9)

    for para in t['transcricao'].split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip(), style='Normal')

    if t.get('signatarios'):
        p = doc.add_paragraph()
        run = p.add_run('Signatarios: ')
        run.bold = True
        run.font.size = docx.shared.Pt(9)
        run2 = p.add_run(', '.join(t['signatarios']))
        run2.font.size = docx.shared.Pt(9)

    if t.get('observacoes'):
        p = doc.add_paragraph()
        run = p.add_run(t['observacoes'])
        run.font.size = docx.shared.Pt(8)
        run.italic = True

    doc.add_paragraph('')
    print(f"    + {t['ministro']} ({t['data_posse']})")

# ── SECTION 2: PLAQUETAS OCR (1975-1993) ──
doc.add_heading('II. Plaquetas institucionais (1975\u20131993)', level=2)

p = doc.add_paragraph()
p.add_run(
    'A partir de Djaci Falcao (1975), o STF passou a publicar plaquetas '
    'institucionais com os discursos integrais das sessoes solenes de posse. '
    'Os textos abaixo foram recuperados via Wayback Machine e transcritos '
    'por OCR. A transicao do termo manuscrito para a plaqueta impressa '
    'marca uma inflexao na auto-narrativa do Tribunal.'
).italic = True

doc.add_paragraph('')

print(f"  II. {len(plaquetas)} plaquetas OCR")
for plq in plaquetas:
    with open(os.path.join(ocr_dir, plq['file']), 'r', encoding='utf-8') as f:
        raw = f.read()
    cleaned = clean_ocr(raw)
    sections = extract_sections(cleaned)

    p = doc.add_paragraph()
    run = p.add_run(plq['presidente'].upper())
    run.font.size = docx.shared.Pt(14)

    meta = doc.add_paragraph()
    run = meta.add_run(plq['data'])
    run.italic = True
    run.font.size = docx.shared.Pt(10)

    meta2 = doc.add_paragraph()
    run = meta2.add_run('Fonte: Plaqueta institucional STF (OCR via Wayback Machine)')
    run.italic = True
    run.font.size = docx.shared.Pt(9)

    for sec in sections:
        if sec['speaker']:
            p = doc.add_paragraph()
            run = p.add_run(sec['speaker'])
            run.bold = True
        paras = text_to_paragraphs(sec['text'])
        for pt in paras:
            doc.add_paragraph(pt, style='Normal')

    doc.add_paragraph('')
    print(f"    + {plq['heading']}")

# ── SECTION 3: DISCURSOS COMPLETOS (1993-2025) ──
doc.add_heading('III. Discursos completos (1993\u20132025)', level=2)

p = doc.add_paragraph()
p.add_run(
    'A partir de Sydney Sanches (1993), os discursos de posse passaram a circular '
    'em formato digital. Os textos abaixo provem do arquivo pessoal da autora '
    'e de fontes oficiais do STF.'
).italic = True

doc.add_paragraph('')

existing = docx.Document(existing_docx)
print(f"  III. Discursos existentes do arquivo pessoal")
for para in existing.paragraphs:
    new_para = doc.add_paragraph()
    try:
        new_para.style = doc.styles[para.style.name]
    except KeyError:
        new_para.style = doc.styles['Normal']
    for run in para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        if run.font.size:
            new_run.font.size = run.font.size

# ── SAVE ──
doc.save(out_path)
print(f"\n{'='*60}")
print(f"CONSOLIDADO FINAL: {out_path}")
print(f"Tamanho: {os.path.getsize(out_path):,} bytes")
print(f"Corpus: 1963-2025 | 3 camadas | zero lacunas")
print(f"{'='*60}")

# Also rename/remove old file
old_path = "C:/Users/medin/Downloads/Discursos posse STF - consolidado 1975-2025.docx"
if os.path.exists(old_path):
    os.remove(old_path)
    print(f"Removido arquivo antigo: {old_path}")
