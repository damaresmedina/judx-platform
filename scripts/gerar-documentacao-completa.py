"""
Gera documentação completa da pesquisa JudX/ICONS
Inclui: diário de achados, instruções de acesso, resultados, inventário
Para caso de emergência — tudo que a pesquisadora precisa sem o Claude
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import psycopg2
import os

JUDX = 'postgresql://postgres:Zb9cHoRww7WxgT0C@db.ejwyguskoiraredinqmb.supabase.co:5432/postgres'
ICONS = 'postgresql://postgres:RHuQvsf4shpsPRjP@db.hetuhkhhppxjliiaerlu.supabase.co:6543/postgres'
OUT = r'C:\Users\medin\Desktop\DOCUMENTACAO_COMPLETA_PESQUISA_JudX_ICONS.docx'

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1a, 0x2a, 0x44)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val) if val is not None else ''
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

def add_note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.italic = True

def query(conn, sql):
    cur = conn.cursor()
    cur.execute(sql)
    rows = cur.fetchall()
    cur.close()
    return rows

# ══════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DOCUMENTAÇÃO COMPLETA DA PESQUISA')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1a, 0x2a, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('JudX · ICONS · Circuitos de Enforcement')
r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Damares Medina')
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Gerado em {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Este documento contém toda a documentação necessária para\ncontinuar a pesquisa independentemente do Claude Code.')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SUMÁRIO
# ══════════════════════════════════════════════════════════

doc.add_heading('Sumário', level=1)
items = [
    'PARTE I — INSTRUÇÕES DE EMERGÊNCIA',
    '  1.1 Como acessar os bancos de dados',
    '  1.2 Onde estão os arquivos',
    '  1.3 Como rodar os scripts',
    '  1.4 Credenciais e acessos',
    'PARTE II — DIÁRIO DE ACHADOS',
    '  2.1 Sessão 22/mar — Fundação',
    '  2.2 Sessão 24/mar — Arquitetura JudX',
    '  2.3 Sessão 25/mar — Limpeza ICONS',
    '  2.4 Sessão 26/mar — Pipeline STF',
    '  2.5 Sessão 27/mar — Análise completa',
    'PARTE III — RESULTADOS EMPÍRICOS',
    '  3.1 STF — Taxa de não-decisão',
    '  3.2 STF — Ambiente virtual e assessorização',
    '  3.3 STF — Coalizões e divergência',
    '  3.4 STJ — Temas repetitivos',
    '  3.5 STJ — Mapa de contaminação',
    '  3.6 ICONS — Ancoragem constitucional',
    'PARTE IV — INVENTÁRIO COMPLETO',
    '  4.1 Estado dos bancos',
    '  4.2 Arquivos locais',
    '  4.3 Scripts disponíveis',
    'PARTE V — COMO CONTINUAR SEM O CLAUDE',
]
for item in items:
    doc.add_paragraph(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PARTE I — INSTRUÇÕES DE EMERGÊNCIA
# ══════════════════════════════════════════════════════════

doc.add_heading('PARTE I — INSTRUÇÕES DE EMERGÊNCIA', level=1)

doc.add_heading('1.1 Como acessar os bancos de dados', level=2)

doc.add_paragraph(
    'Os dados da pesquisa estão em dois bancos PostgreSQL no Supabase. '
    'Você pode acessá-los pelo painel web do Supabase (sem programação).'
)

doc.add_heading('Banco JudX (dados do STF e STJ)', level=3)
doc.add_paragraph('1. Abra o navegador e vá para: https://supabase.com/dashboard')
doc.add_paragraph('2. Faça login com sua conta Google')
doc.add_paragraph('3. Selecione o projeto "judx-platform" (ID: ejwyguskoiraredinqmb)')
doc.add_paragraph('4. No menu lateral, clique em "Table Editor" para ver as tabelas')
doc.add_paragraph('5. Clique em qualquer tabela para ver os dados')
doc.add_paragraph('6. Use "SQL Editor" no menu lateral para rodar consultas')

doc.add_heading('Banco ICONS (ancoragem constitucional)', level=3)
doc.add_paragraph('1. No mesmo painel Supabase, troque para o projeto "icons" (ID: hetuhkhhppxjliiaerlu)')
doc.add_paragraph('2. Mesmo procedimento — Table Editor ou SQL Editor')

add_note('Dica: no SQL Editor, copie e cole as queries que estão na Parte III deste documento para reproduzir qualquer resultado.')

doc.add_heading('1.2 Onde estão os arquivos', level=2)

add_table(
    ['O que', 'Onde', 'Para que serve'],
    [
        ['Projeto JudX (código)', r'C:\Users\medin\projetos\judx-platform', 'Scripts, skills, protocolo'],
        ['Scripts de extração', r'C:\Users\medin\projetos\judx-platform\scripts', '7 scripts de pipeline'],
        ['Backups em Excel', r'C:\Users\medin\Desktop\backup_judx', 'Cópia local dos bancos + achados'],
        ['Papers Bicocca', r'C:\Users\medin\Desktop\bicocca milano', 'Papers acadêmicos e subsidiários'],
        ['Paper atual', r'C:\Users\medin\Downloads\CircuitosEnforcement_COMPLETO (1).docx', 'Draft do Circuitos de Enforcement'],
        ['Manual voz autoral', r'C:\Users\medin\Desktop\infoprodutos\manual da voz autoral dm.docx', 'Referência de estilo'],
        ['FIESP pesquisa', r'C:\Users\medin\Downloads\file-20230622145000-fiesp-*.pdf', 'Obstáculos ao crescimento'],
        ['Teto Decorativo', r'C:\Users\medin\Desktop\bicocca milano\subsidios\TETO-DECORATIVO-*.pdf', 'Honorários sucumbência'],
        ['Projeto ICONS', r'C:\projetos\icons', 'Banco ICONS local'],
        ['Site ICONS (deploy)', r'C:\projetos\icons-cartografia', 'HTML para icons.org.br'],
        ['Dados brutos STF', r'C:\projetos\judx\STF', 'Planilhas Excel originais'],
        ['Curso IDP', r'C:\Users\medin\Desktop\scc_idp', 'Material do curso SCC'],
    ]
)

doc.add_heading('1.3 Como rodar os scripts', level=2)

doc.add_paragraph('Todos os scripts rodam pelo terminal (prompt de comando):')
doc.add_paragraph('1. Abra o terminal: tecle Windows + R, digite "cmd", Enter')
doc.add_paragraph('2. Navegue até a pasta: cd C:\\Users\\medin\\projetos\\judx-platform')
doc.add_paragraph('3. Para scripts JavaScript (.mjs): node scripts/nome-do-script.mjs')
doc.add_paragraph('4. Para scripts Python (.py): python scripts/nome-do-script.py')

add_table(
    ['Script', 'O que faz', 'Como rodar'],
    [
        ['bom-dia.mjs', 'Diagnóstico completo de estado', 'node scripts/bom-dia.mjs'],
        ['inventario-completo.py', 'Gera Excel com inventário', 'python scripts/inventario-completo.py'],
        ['backup-banco.py', 'Exporta bancos para Excel', 'python scripts/backup-banco.py'],
        ['gerar-relatorio.py', 'Gera relatório Word de achados', 'python scripts/gerar-relatorio.py'],
        ['run-stf-pipeline-fast.mjs', 'Normaliza decisões STF', 'node scripts/run-stf-pipeline-fast.mjs'],
        ['fetch-stf-partes-safe.mjs', 'Extrai partes do portal STF', 'node scripts/fetch-stf-partes-safe.mjs'],
        ['fetch-stj-temas.mjs', 'Extrai temas repetitivos STJ', 'node scripts/fetch-stj-temas.mjs'],
    ]
)

doc.add_heading('1.4 Credenciais e acessos', level=2)

add_table(
    ['Serviço', 'Credencial'],
    [
        ['Supabase JudX', 'Senha banco: Zb9cHoRww7WxgT0C'],
        ['Supabase ICONS', 'Senha banco: RHuQvsf4shpsPRjP'],
        ['GitHub', 'damaresmedina (auth via gh CLI — já configurado)'],
        ['Vercel', 'damaresmedinas-projects (login via npx vercel)'],
        ['Datajud CNJ API Key', 'cDZHYzlZa0JadVREZDJCendQbXY6SkJlTzNjLV9TRENyQk1RdnFKZGRQdw=='],
    ]
)

add_note('IMPORTANTE: guarde estas credenciais em local seguro. Sem elas não é possível acessar os bancos.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PARTE II — DIÁRIO DE ACHADOS
# ══════════════════════════════════════════════════════════

doc.add_heading('PARTE II — DIÁRIO DE ACHADOS', level=1)

doc.add_paragraph(
    'Este diário registra cronologicamente tudo o que foi descoberto, construído e decidido em cada sessão '
    'de trabalho. Cada achado inclui a fonte dos dados, o método usado e as limitações conhecidas.'
)

# Read the diary file
diary_path = r'C:\Users\medin\projetos\judx-platform\DIARIO_ACHADOS.md'
with open(diary_path, 'r', encoding='utf-8') as f:
    diary = f.read()

# Parse and add sections
sections = diary.split('## ')
for section in sections[1:]:  # skip header
    lines = section.strip().split('\n')
    title = lines[0].strip()
    doc.add_heading(title, level=2)

    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PARTE III — RESULTADOS EMPÍRICOS (com queries e dados live)
# ══════════════════════════════════════════════════════════

doc.add_heading('PARTE III — RESULTADOS EMPÍRICOS', level=1)

print('Conectando ao banco para dados live...')
conn = psycopg2.connect(JUDX, sslmode='require')

# 3.1 Não-decisão
doc.add_heading('3.1 STF — Taxa de não-decisão', level=2)
doc.add_paragraph(
    'O dado mais revelador do corpus JudX: mais de 79% das decisões do STF não apreciam o mérito da causa.'
)

rows = query(conn, """
SELECT descricao_andamento, COUNT(*) as n,
  ROUND(COUNT(*)*100.0/SUM(COUNT(*)) OVER(),1) as pct
FROM stf_decisoes GROUP BY 1 ORDER BY 2 DESC LIMIT 15
""")
add_table(['Resultado', 'Total', '%'], [[r[0] or '?', f'{r[1]:,}', f'{r[2]}%'] for r in rows])
add_note('Fonte: corpus JudX, 169.851 decisões, Corte Aberta STF 1988-2026')

# 3.2 Virtual
doc.add_heading('3.2 STF — Ambiente virtual e assessorização', level=2)

rows = query(conn, """
SELECT orgao_julgador, COUNT(*) as total,
  SUM(CASE WHEN observacao_andamento ILIKE '%unanimidade%' THEN 1 ELSE 0 END) as unanimes,
  ROUND(SUM(CASE WHEN observacao_andamento ILIKE '%unanimidade%' THEN 1 ELSE 0 END)*100.0/NULLIF(COUNT(*),0),1) as pct
FROM stf_decisoes WHERE tipo_decisao='COLEGIADA' AND decisoes_virtual=true
GROUP BY 1 HAVING COUNT(*)>50 ORDER BY 2 DESC
""")
add_table(['Órgão', 'Total', 'Unânimes', '% Unânime'],
    [[r[0], f'{r[1]:,}', f'{r[2]:,}', f'{r[3]}%'] for r in rows])

doc.add_paragraph('')
doc.add_paragraph('Semanas de maior volume:')
rows = query(conn, """
SELECT DATE_TRUNC('week',TO_DATE(data_decisao,'DD/MM/YYYY'))::date, COUNT(*),
  COUNT(DISTINCT relator_decisao), ROUND(COUNT(*)*1.0/NULLIF(COUNT(DISTINCT relator_decisao),0),1)
FROM stf_decisoes WHERE decisoes_virtual=true AND tipo_decisao='COLEGIADA' AND data_decisao ~ '^\\d{2}/\\d{2}/\\d{4}'
GROUP BY 1 ORDER BY 2 DESC LIMIT 10
""")
add_table(['Semana', 'Processos', 'Relatores', 'Média/relator'],
    [[str(r[0]), f'{r[1]:,}', str(r[2]), str(r[3])] for r in rows])
add_note('Média/relator exclui "MINISTRO PRESIDENTE" e "VICE-PRESIDENTE" na análise detalhada (ver diário).')

# 3.3 Coalizões
doc.add_heading('3.3 STF — Coalizões e divergência', level=2)

doc.add_paragraph('Série histórica da divergência (% de acórdãos com "vencido" no texto):')
rows = query(conn, """
SELECT SUBSTRING(data_decisao FROM '\\d{4}$')::int as ano, COUNT(*),
  SUM(CASE WHEN observacao_andamento ILIKE '%vencido%' THEN 1 ELSE 0 END),
  ROUND(SUM(CASE WHEN observacao_andamento ILIKE '%vencido%' THEN 1 ELSE 0 END)*100.0/NULLIF(COUNT(*),0),1)
FROM stf_decisoes WHERE tipo_decisao='COLEGIADA' AND decisoes_virtual=true
  AND data_decisao ~ '\\d{2}/\\d{2}/\\d{4}' AND SUBSTRING(data_decisao FROM '\\d{4}$')::int >= 2016
GROUP BY 1 ORDER BY 1
""")
add_table(['Ano', 'Acórdãos', 'Com divergência', '%'],
    [[str(r[0]), f'{r[1]:,}', f'{r[2]:,}', f'{r[3]}%'] for r in rows])

doc.add_paragraph('')
doc.add_paragraph('Ministros explicitamente vencidos (top 15):')
rows = query(conn, """
SELECT (regexp_matches(observacao_andamento,
  'vencid[ao]s?\\s+(?:o|a|os|as)\\s+Ministr[ao]s?\\s+([^,.;\\u2013\\u2014]+)','i'))[1], COUNT(*)
FROM stf_decisoes WHERE tipo_decisao='COLEGIADA' AND observacao_andamento ~* 'vencid[ao]'
GROUP BY 1 ORDER BY 2 DESC LIMIT 15
""")
add_table(['Ministro vencido', 'Vezes'], [[r[0].strip(), f'{r[1]:,}'] for r in rows])

doc.add_paragraph('')
doc.add_paragraph('Contrafactual Marco Aurélio:')
rows = query(conn, """
SELECT SUBSTRING(data_decisao FROM '\\d{4}$')::int,COUNT(*),
  ROUND(SUM(CASE WHEN observacao_andamento ILIKE '%vencido%' THEN 1 ELSE 0 END)*100.0/NULLIF(COUNT(*),0),1),
  ROUND(SUM(CASE WHEN observacao_andamento ILIKE '%vencido%' AND relator_decisao NOT ILIKE '%MARCO%' THEN 1 ELSE 0 END)*100.0/
  NULLIF(SUM(CASE WHEN relator_decisao NOT ILIKE '%MARCO%' THEN 1 ELSE 0 END),0),1)
FROM stf_decisoes WHERE tipo_decisao='COLEGIADA' AND decisoes_virtual=true
  AND data_decisao ~ '\\d{2}/\\d{2}/\\d{4}' AND SUBSTRING(data_decisao FROM '\\d{4}$')::int BETWEEN 2018 AND 2025
GROUP BY 1 ORDER BY 1
""")
add_table(['Ano', 'Total', '% com MA', '% sem MA'],
    [[str(r[0]), f'{r[1]:,}', f'{r[2]}%', f'{r[3]}%'] for r in rows])
add_note('Delta máximo: 1,1pp (2020-2021). Marco Aurélio não explica a anomalia.')

# 3.4 STJ Temas
doc.add_heading('3.4 STJ — Temas repetitivos', level=2)

rows = query(conn, """
SELECT situacao, COUNT(*) FROM stj_temas GROUP BY 1 ORDER BY 2 DESC
""")
add_table(['Situação', 'Temas'], [[r[0], str(r[1])] for r in rows])

doc.add_paragraph('')
rows = query(conn, """
SELECT ramo_direito, COUNT(*), ROUND(AVG(data_julgamento-data_afetacao))
FROM stj_temas WHERE data_afetacao IS NOT NULL AND data_julgamento IS NOT NULL
GROUP BY 1 ORDER BY 3 DESC
""")
add_table(['Ramo', 'Temas', 'Dias médio'],
    [[r[0], str(r[1]), str(r[2]) if r[2] else '—'] for r in rows])

# 3.5 Contaminação
doc.add_heading('3.5 STJ — Mapa de contaminação', level=2)

rows = query(conn, """
SELECT tribunal_origem, COUNT(*), ROUND(COUNT(*)*100.0/(SELECT COUNT(*) FROM stj_processos_semente WHERE tribunal_origem IS NOT NULL),1)
FROM stj_processos_semente WHERE tribunal_origem IS NOT NULL GROUP BY 1 ORDER BY 2 DESC LIMIT 10
""")
add_table(['Tribunal de origem', 'Sementes', '%'], [[r[0], str(r[1]), f'{r[2]}%'] for r in rows])

doc.add_paragraph('')
rows = query(conn, """
SELECT relator, COUNT(*), COUNT(*) FILTER (WHERE situacao='transito_em_julgado'),
  COUNT(*) FILTER (WHERE situacao='afetado')
FROM stj_temas WHERE relator IS NOT NULL GROUP BY 1 ORDER BY 2 DESC LIMIT 10
""")
add_table(['Relator', 'Temas', 'Trânsito', 'Pendentes'],
    [[r[0], str(r[1]), str(r[2]), str(r[3])] for r in rows])

# 3.6 ICONS
doc.add_heading('3.6 ICONS — Ancoragem constitucional', level=2)

conn_icons = psycopg2.connect(ICONS, sslmode='require')
rows = query(conn_icons, """
SELECT o.slug, COUNT(DISTINCT e.source_id) as n
FROM objects o JOIN edges e ON e.target_id = o.id
WHERE e.type_slug='ancora_normativa' AND o.type_slug='artigo'
GROUP BY o.id, o.slug ORDER BY n DESC LIMIT 15
""")
add_table(['Dispositivo', 'Decisões vinculadas'], [[r[0], f'{r[1]:,}'] for r in rows])
add_note('7.766 edges ancora_normativa. 100% source_id são registro_jurisprudencial.')
conn_icons.close()

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PARTE IV — INVENTÁRIO
# ══════════════════════════════════════════════════════════

doc.add_heading('PARTE IV — INVENTÁRIO COMPLETO', level=1)

doc.add_heading('4.1 Estado dos bancos', level=2)

tables_info = [
    ('stf_decisoes', 'Decisões raw do STF (Corte Aberta)'),
    ('judx_case', 'Processos normalizados STF'),
    ('judx_decision', 'Decisões normalizadas STF'),
    ('stf_partes', 'Partes processuais STF'),
    ('stj_temas', 'Temas repetitivos STJ'),
    ('stj_processos_semente', 'Processos-semente STJ'),
    ('stj_contramostra', 'Contramostra STJ'),
    ('judx_court', 'Tribunais'),
    ('judx_judge', 'Ministros/juízes'),
    ('judx_subject', 'Assuntos'),
    ('judx_procedural_class', 'Classes processuais'),
]

inv_rows = []
for tname, desc in tables_info:
    try:
        cnt = query(conn, f"SELECT COUNT(*) FROM {tname}")[0][0]
        inv_rows.append([tname, desc, f'{cnt:,}'])
    except:
        inv_rows.append([tname, desc, 'ERRO'])
        conn.rollback()

add_table(['Tabela', 'Descrição', 'Registros'], inv_rows)

doc.add_heading('4.2 Sites no ar', level=2)
add_table(['Site', 'URL'], [
    ['ICONS', 'https://icons.org.br'],
    ['JudX', 'https://judx-platform.vercel.app'],
])

doc.add_heading('4.3 Backups em Excel', level=2)
doc.add_paragraph(f'Pasta: C:\\Users\\medin\\Desktop\\backup_judx\\')
backup_dir = r'C:\Users\medin\Desktop\backup_judx'
if os.path.exists(backup_dir):
    files = sorted(os.listdir(backup_dir))
    bk_rows = []
    for f in files:
        fp = os.path.join(backup_dir, f)
        sz = os.path.getsize(fp) / (1024*1024)
        bk_rows.append([f, f'{sz:.1f} MB'])
    add_table(['Arquivo', 'Tamanho'], bk_rows)

conn.close()

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PARTE V — COMO CONTINUAR SEM O CLAUDE
# ══════════════════════════════════════════════════════════

doc.add_heading('PARTE V — COMO CONTINUAR SEM O CLAUDE', level=1)

doc.add_heading('Se precisar acessar os dados', level=2)
doc.add_paragraph('1. Abra os arquivos Excel em Desktop\\backup_judx\\ — são cópias completas dos bancos')
doc.add_paragraph('2. O arquivo INVENTARIO_COMPLETO tem 9 abas com tudo organizado')
doc.add_paragraph('3. O arquivo achados_completos tem todas as análises em abas separadas')

doc.add_heading('Se precisar acessar o banco online', level=2)
doc.add_paragraph('1. Vá para https://supabase.com/dashboard')
doc.add_paragraph('2. Login → projeto judx-platform → SQL Editor')
doc.add_paragraph('3. As queries estão neste documento (Parte III) — copie e cole')

doc.add_heading('Se precisar de um novo Claude Code', level=2)
doc.add_paragraph('1. Abra o Claude Code no terminal')
doc.add_paragraph('2. Navegue para C:\\Users\\medin\\projetos\\judx-platform')
doc.add_paragraph('3. O Claude vai ler automaticamente o CLAUDE.md e o STATUS.md')
doc.add_paragraph('4. Diga "rode o bom-dia" — ele faz o diagnóstico completo')
doc.add_paragraph('5. Todas as memórias, skills e instruções estão no projeto')

doc.add_heading('Se precisar citar dados no paper', level=2)
doc.add_paragraph('Todos os números confirmados estão na Parte III deste documento com:')
doc.add_paragraph('• O dado exato')
doc.add_paragraph('• A fonte (corpus, tabela, query)')
doc.add_paragraph('• A limitação (proxy usada, amostra, período)')
doc.add_paragraph('')
doc.add_paragraph('Use "Fonte: a autora, a partir de dados do STF" quando citar dados do corpus JudX.')
doc.add_paragraph('Use a fonte específica quando citar dados externos (FIESP, CNJ, Transparência Brasil).')

doc.add_heading('Contatos de emergência técnica', level=2)
doc.add_paragraph('• GitHub: https://github.com/damaresmedina/judx-platform (código + histórico)')
doc.add_paragraph('• Vercel: https://vercel.com/damaresmedinas-projects (deploys)')
doc.add_paragraph('• Supabase: https://supabase.com/dashboard (bancos)')

# Footer
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Documento gerado automaticamente em {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PROJUS · Instituto Constituição Aberta — ICONS · JudX Platform')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print(f'\nDocumentação salva: {OUT}')
