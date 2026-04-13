import docx, sys, os, json
sys.stdout.reconfigure(encoding='utf-8')

termos = [
    {
        "arquivo": "006 (2).pdf",
        "ministro": "Luiz Gallotti",
        "data_posse": "01/12/1966",
        "presidente_sessao": "Ribeiro da Costa",
        "procurador_geral": "nao identificado",
        "signatarios": ["Luiz Gallotti"],
        "transcricao": "Termo de posse do Sr. Ministro Luiz Gallotti no cargo de Presidente do Supremo Tribunal Federal.\n\nAos primeiro dias do mes de Dezembro de 1966, na Sala de Sessoes do Supremo Tribunal Federal, presentes os Exmos. Srs. Ministros Luiz Gallotti, Presidente, dos Srs. Ministros, dos quais os abaixo assinam o presente termo e do Sr. Procurador Geral da Republica, de altas autoridades, compareceu o Sr. Ministro Luiz Gallotti que em vista de ter sido eleito para o cargo de mandato do Sr. Ministro Presidente, e em decorrencia de eleicao havida no dia sete do mesmo mes, eleicao valida para o bienio 67/68, tendo o Sr. Ministro Gallotti, depois de prestar o compromisso de bem cumprir os deveres do cargo, de acordo com as leis da Republica, assumiu suas funcoes. E para constar, lavrei o presente termo, que vai assinado pelo Sr. Ministro Presidente e pelo empossado e pelos demais Srs. Ministros. E eu, Bacharel Hugo Mosca, escrevi e rubrico o presente termo.\n\nLuiz Gallotti",
        "observacoes": "Manuscrito de dificil leitura. Caligrafia cursiva informal. Pagina 120 do livro de termos."
    },
    {
        "arquivo": "006 (1).pdf",
        "ministro": "Antonio Goncalves de Oliveira",
        "data_posse": "12/12/1968",
        "presidente_sessao": "Luiz Gallotti",
        "procurador_geral": "Procurador-Geral da Republica",
        "signatarios": ["Luiz Gallotti", "Antonio Goncalves de Oliveira", "Oswaldo Trigueiro de Albuquerque Melo", "Victor Nunes Leal", "Evandro Lins e Silva", "Adaucto Cardoso"],
        "transcricao": "Termo de posse do Exmo. Sr. Ministro Antonio Goncalves de Oliveira no cargo de Presidente do Supremo Tribunal Federal.\n\nAos doze dias do mes de Dezembro de 1968, na Sala de Sessoes do Supremo Tribunal Federal, presentes os Exmos. Srs. Ministros Luiz Gallotti, Presidente, dos Srs. Ministros que assinam o presente termo, do Exmo. Sr. Procurador Geral da Republica, de altas autoridades, compareceu o Exmo. Sr. Ministro Antonio Goncalves de Oliveira que, em vista de ter sido eleito, em concordancia de eleicao realizada no dia 11 do mesmo mes, foi eleito Presidente do Supremo Tribunal Federal para o bienio 69-70, tendo o Exmo. Sr. Ministro Antonio Goncalves de Oliveira, depois de prestar o compromisso de bem cumprir os deveres do cargo, de acordo com as leis da Republica, assumiu suas funcoes. E, para constar, lavrei o presente termo que vai assinado pelo Exmo. Sr. Ministro Presidente, pelo empossado e pelos demais Srs. Ministros. E eu, Bacharel Hugo Mosca, assino e rubrico o presente termo.",
        "observacoes": "Manuscrito cursivo informal. Pagina 131 do livro de termos."
    },
    {
        "arquivo": "008.pdf",
        "ministro": "Oswaldo Trigueiro de Albuquerque Melo",
        "data_posse": "10/02/1969",
        "presidente_sessao": "Luiz Gallotti",
        "procurador_geral": "Procurador-Geral da Republica",
        "signatarios": ["Oswaldo Trigueiro de Albuquerque Melo", "Adaucto Nogueira", "Aliomar Baleeiro", "Eloy da Rocha", "Djaci Falcao", "Barros Monteiro", "Amaral Santos", "Thompson Flores"],
        "transcricao": "Termo de posse do Exmo. Sr. Ministro Oswaldo Trigueiro de Albuquerque Melo -- na Presidencia deste Tribunal.\n\nAos dez dias do mes de fevereiro do ano de mil novecentos e sessenta e nove, na Sala de Sessoes do Supremo Tribunal Federal, sob a Presidencia do Exmo. Sr. Ministro Luiz Gallotti, na presenca dos Exmos. Srs. Ministros que assinam este termo e do Exmo. Sr. Dr. Procurador Geral da Republica, compareceu o Exmo. Sr. Ministro Oswaldo Trigueiro de Albuquerque Melo, que, em virtude da eleicao realizada a 6 do corrente, na qual foi eleito Presidente por um bienio, a partir desta data, tomou posse e entrou no exercicio do cargo, depois de prestar o compromisso de bem cumprir os deveres do mesmo, de acordo com as leis da Republica. Em seguida, o Exmo. Sr. Ministro-Presidente eleito, assumiu as funcoes. E para constar, determinou o Exmo. Sr. Ministro Presidente, a lavratura do presente termo, que vai assinado por Sua Excia., pelo empossado, e pelos demais Senhores Ministros desta Corte. E eu, Joely Cruz, Diretor-Geral da Secretaria, o escrevi e assinei.",
        "observacoes": "Manuscrito em caligrafia formal. Paginas 132-133 do livro de termos. Escrivao: Joely Cruz."
    },
    {
        "arquivo": "007.pdf",
        "ministro": "Aliomar de Andrade Baleeiro",
        "data_posse": "10/02/1971",
        "presidente_sessao": "Oswaldo Trigueiro de Albuquerque Melo",
        "procurador_geral": "Francisco Manoel Xavier de Albuquerque",
        "signatarios": ["Oswaldo Trigueiro", "Aliomar Baleeiro", "Adaucto Nogueira", "Djaci Falcao", "Barros Monteiro", "Carlos Campos", "Bilac Pinto"],
        "transcricao": "Termo de posse do Excelentissimo Senhor Ministro Aliomar de Andrade Baleeiro, no cargo de Presidente do Supremo Tribunal Federal.\n\nAos dez dias do mes de fevereiro de mil novecentos e setenta e um, perante o Supremo Tribunal Federal, reunido em sessao solene, sob a Presidencia do Excelentissimo Senhor Ministro Oswaldo Trigueiro de Albuquerque Melo, e presente o Excelentissimo Senhor Doutor Francisco Manoel Xavier de Albuquerque, Procurador Geral da Republica, compareceu o Excelentissimo Senhor Ministro Aliomar de Andrade Baleeiro, eleito em oito do mesmo mes, Presidente do Supremo Tribunal Federal, para o bienio mil novecentos e setenta e um, mil novecentos e setenta e dois, em virtude do termino do mandato do Excelentissimo Senhor Ministro Oswaldo Trigueiro de Albuquerque Melo, o qual, depois do compromisso legal de bem e fielmente cumprir os deveres do cargo, assumiu o exercicio de suas funcoes. E para constar, lavrou-se o presente termo que vai assinado pelo Excelentissimo Senhor Ministro Presidente, pelo empossado, pelos demais Ministros presentes, pelo Doutor Procurador Geral da Republica, e por mim, Diretor Geral da Secretaria que o escrevi.",
        "observacoes": "Caligrafia ornamental muito elegante."
    },
    {
        "arquivo": "006.pdf",
        "ministro": "Eloy Jose da Rocha",
        "data_posse": "09/02/1973",
        "presidente_sessao": "Aliomar de Andrade Baleeiro",
        "procurador_geral": "Jose Carlos Moreira Alves",
        "signatarios": ["Aliomar Baleeiro", "Eloy da Rocha", "Oswaldo Trigueiro", "Djaci Falcao", "Carlos Campos", "Antonio Neder", "Cordeiro Guerra", "Rodrigues Alckmin", "Xavier de Albuquerque", "Bilac Pinto"],
        "transcricao": "Termo de posse do Excelentissimo Senhor Ministro Eloy Jose da Rocha, no cargo de Presidente do Supremo Tribunal Federal.\n\nAos nove dias do mes de fevereiro de mil novecentos e setenta e tres, perante o Supremo Tribunal Federal, reunido em Sessao Solene, sob a Presidencia do Excelentissimo Senhor Ministro Aliomar de Andrade Baleeiro, e presente o Excelentissimo Senhor Doutor Jose Carlos Moreira Alves, Procurador Geral da Republica, compareceu o Excelentissimo Senhor Ministro Eloy Jose da Rocha, eleito em sete do mesmo mes, Presidente do Supremo Tribunal Federal, para o bienio mil novecentos e setenta e tres, mil novecentos e setenta e quatro, em virtude do termino do mandato do Excelentissimo Senhor Ministro Aliomar de Andrade Baleeiro; o qual, depois do compromisso legal de bem e fielmente cumprir os deveres do cargo, assumiu o exercicio de suas funcoes. E, para constar, lavrou-se o presente termo que vai assinado pelo Excelentissimo Senhor Ministro Presidente, pelo empossado, pelos demais Ministros presentes, pelo Doutor Procurador Geral da Republica e por mim, Jaqueline Juris, Diretor Geral da Secretaria.",
        "observacoes": "Caligrafia ornamental elegante. Paginas 140-141 do livro de termos."
    },
    {
        "arquivo": "005.pdf",
        "ministro": "Antonio Neder",
        "data_posse": "14/02/1979",
        "presidente_sessao": "Carlos Thompson Flores",
        "procurador_geral": "Henrique Fonseca de Araujo",
        "signatarios": ["Carlos Thompson Flores", "Antonio Neder", "Djaci Falcao", "Leitao de Abreu", "Moreira Alves", "Decio Miranda", "Soares Munoz", "Henrique Fonseca de Araujo"],
        "transcricao": "Termo de Posse do Excelentissimo Senhor Ministro Antonio Neder, no cargo de Presidente do Supremo Tribunal Federal.\n\nAos quatorze dias do mes de fevereiro do ano de mil novecentos e setenta e nove, perante o Supremo Tribunal Federal, reunido em Sessao Solene, sob a presidencia do Excelentissimo Senhor Ministro Carlos Thompson Flores, presente o Senhor Procurador-Geral da Republica, Professor Henrique Fonseca de Araujo, compareceu o Excelentissimo Senhor Ministro Antonio Neder, eleito Presidente do Supremo Tribunal Federal, em sessao de treze de dezembro de mil novecentos e setenta e oito, para o bienio mil novecentos e setenta e nove -- mil novecentos e oitenta e um, em virtude do termino do mandato do Excelentissimo Senhor Ministro Carlos Thompson Flores, o qual, depois de prestar o compromisso legal de bem e fielmente cumprir os deveres do cargo, assumiu o exercicio de suas funcoes. E, para constar, lavrou-se o presente Termo, que vai assinado pelo Excelentissimo Senhor Presidente, pelo empossado, pelos demais Ministros presentes, pelo Procurador-Geral da Republica e por mim, Tito Di Ramos, Diretor-Geral da Secretaria.",
        "observacoes": "Manuscrito formal em papel pautado."
    },
    {
        "arquivo": "008 (1).pdf",
        "ministro": "Alvaro Moutinho Ribeiro da Costa",
        "data_posse": "11/12/1963",
        "presidente_sessao": "Antonio Carlos Lafayette de Andrada",
        "procurador_geral": "Candido de Oliveira Neto",
        "signatarios": ["Lafayette de Andrada", "Ribeiro da Costa", "Luiz Gallotti", "Candido Motta Filho", "Hahnemann Guimaraes", "Vilas Boas", "Victor Nunes Leal", "Pedro Chaves", "Goncalves de Oliveira", "Evandro Lins"],
        "tipo": "Diario da Justica impresso",
        "transcricao": "SUPREMO TRIBUNAL FEDERAL\nATA DA TRIGESIMA NONA SESSAO, EM 11 DE DEZEMBRO DE 1963\n\nPresidencia do Exmo. Sr. Ministro Antonio Carlos Lafayette de Andrada. Procurador-Geral da Republica, o Exmo. Sr. Dr. Candido de Oliveira Neto. Secretario, o Dr. Hugo Mosca, Vice-Diretor Geral.\n\nAs treze horas, abriu-se a sessao, achando-se presentes os Exmos. Senhores Ministros Evandro Lins, Pedro Chaves, Victor Nunes Leal, Goncalves de Oliveira, Vilas Boas, Candido Motta Filho, Luiz Gallotti, Hahnemann Guimaraes e Ribeiro da Costa.\n\nPOSSE DO EXMO. SR. MINISTRO A. M. RIBEIRO DA COSTA NA PRESIDENCIA DO SUPREMO TRIBUNAL FEDERAL\n\nAprovada a ata, o Exmo. Sr. Ministro A. C. Lafayette de Andrada convidou os Exmos. Srs. Ministros Luiz Gallotti, Candido Motta Filho e Vilas Boas para introduzirem no recinto o Exmo. Sr. Ministro A. M. Ribeiro da Costa, eleito presidente do Supremo Tribunal Federal para o bienio 64-65.\n\n[Discursos integrais publicados no DJ 12/12/1963, p. 4364-4366]",
        "discurso_disponivel": True,
        "observacoes": "FONTE PRIMARIA COMPLETA. DJ 12/12/1963, Ano XXXVIII, No. 235. 4 paginas com discursos integrais."
    },
]

# Save JSON
out_json = "C:/Users/medin/Desktop/backup_judx/resultados/termos_posse_stf_1963_1979.json"
with open(out_json, 'w', encoding='utf-8') as f:
    json.dump(termos, f, ensure_ascii=False, indent=2)
print(f"JSON: {out_json}")

# Build docx
doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = docx.shared.Pt(11)

p = doc.add_paragraph()
run = p.add_run('TERMOS DE POSSE NA PRESIDENCIA DO STF (1963-1979)')
run.font.size = docx.shared.Pt(16)
run.bold = True
p = doc.add_paragraph()
run = p.add_run('Transcritos dos manuscritos originais e do Diario da Justica')
run.italic = True
doc.add_paragraph('')

termos_sorted = sorted(termos, key=lambda t: t['data_posse'].split('/')[::-1])

for t in termos_sorted:
    doc.add_heading(f"{t['ministro']} ({t['data_posse']})", level=3)
    meta = doc.add_paragraph()
    meta.add_run('Sessao presidida por: ').bold = True
    meta.add_run(f"{t['presidente_sessao']}\n")
    meta.add_run('Procurador-Geral: ').bold = True
    meta.add_run(f"{t['procurador_geral']}\n")
    meta.add_run('Fonte: ').bold = True
    meta.add_run(f"{t.get('tipo', 'Manuscrito original')}\n")
    meta.add_run('Obs: ').bold = True
    meta.add_run(t['observacoes'])
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('TRANSCRICAO:').bold = True
    for para in t['transcricao'].split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip(), style='Normal')
    if t.get('signatarios'):
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('Signatarios: ').bold = True
        p.add_run(', '.join(t['signatarios']))
    doc.add_paragraph('')

termos_path = "C:/Users/medin/Downloads/Termos de posse STF 1963-1979 - transcritos.docx"
doc.save(termos_path)
print(f"DOCX: {termos_path} ({os.path.getsize(termos_path):,} bytes)")
print(f"\n7 termos transcritos, ordem cronologica 1963-1979")
for t in termos_sorted:
    print(f"  {t['data_posse']} - {t['ministro']}")
