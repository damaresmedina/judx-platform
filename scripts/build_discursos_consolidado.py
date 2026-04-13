import docx, sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

ocr_dir = "C:/Users/medin/Desktop/backup_judx/resultados/plaquetas_stf/ocr_txt"

def clean_ocr(text):
    text = re.sub(r'--- Pagina \d+ ---', '', text)
    text = re.sub(r'^#.*\n', '', text)
    # zero -> o
    text = text.replace('a0 ', 'ao ')
    text = text.replace(' 0 ', ' o ')
    text = re.sub(r'\b0(?=[a-z\u00e0-\u00ff])', 'o', text)
    text = re.sub(r'(?<=[a-z\u00e0-\u00ff])0\b', 'o', text)
    # broken hyphens
    text = re.sub(r'(\w)\s*-\s*\n\s*(\w)', r'\1\2', text)
    text = re.sub(r'(\w)\s+-\s+(\w)', r'\1\2', text)
    # OCR marks
    text = text.replace('\u20ac', 'e')
    text = text.replace('_', '.')
    text = text.replace('<', '').replace('>', '')
    # stray punctuation lines
    text = re.sub(r'\n\s*[;:,.\s]\s*\n', '\n', text)
    # page numbers alone
    text = re.sub(r'\n\s*\d{1,3}\s*\n', '\n', text)
    # whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\s+([.;,!?:])', r'\1', text)
    return text.strip()

def extract_sections(text):
    pattern = r'(Palavras\s+d[eo]\s+(?:Senhor\s+)?Ministro\s+[^\n]+)'
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    sections = []
    current_speaker = None
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r'Palavras\s+d[eo]', part, re.IGNORECASE):
            current_speaker = part.strip()
        else:
            if 'POSSE NA PRESID\u00caNCIA' in part and len(part) < 400:
                continue
            if 'SOLENIDADE DE POSSE' in part and len(part) < 500:
                continue
            if 'REP\u00daBLICA FEDERATIVA' in part and len(part) < 500:
                continue
            sections.append({'speaker': current_speaker, 'text': part})
    return sections

def text_to_paragraphs(text):
    raw_paras = re.split(r'\n\n+', text)
    paras = []
    for p in raw_paras:
        p = p.strip()
        p = re.sub(r'\n', ' ', p)
        p = re.sub(r'\s+', ' ', p)
        if len(p) > 10:
            paras.append(p)
    return paras

plaquetas = [
    {"heading": "Djaci Falc\u00e3o (1975-1977)",
     "file": "Plaqueta_Possepresidencial_DjaciFalcao_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de fevereiro de 1975",
     "presidente": "Djaci Alves Falc\u00e3o",
     "vice": "Thompson Flores"},
    {"heading": "Thompson Flores (1977-1979)",
     "file": "Plaqueta_Possepresidencial_ThompsonFlores_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de fevereiro de 1977",
     "presidente": "Carlos Thompson Flores",
     "vice": "Olavo Bilac Pinto"},
    {"heading": "Xavier de Albuquerque (1981-1983)",
     "file": "Plaqueta_Possepresidencial_XavierdeAlbuquerquearrumar_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 16 de fevereiro de 1981",
     "presidente": "Francisco Manoel Xavier de Albuquerque",
     "vice": "Jo\u00e3o Leit\u00e3o de Abreu"},
    {"heading": "Cordeiro Guerra (1983-1985)",
     "file": "Plaqueta_Possepresidencial_CordeiroGuerra_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 21 de fevereiro de 1983",
     "presidente": "Jo\u00e3o Baptista Cordeiro Guerra",
     "vice": "Jos\u00e9 Carlos Moreira Alves"},
    {"heading": "Moreira Alves (1985-1987)",
     "file": "Plaqueta_Possepresidencial_MoreiraAlves_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 25 de fevereiro de 1985",
     "presidente": "Jos\u00e9 Carlos Moreira Alves",
     "vice": "D\u00e9cio Miranda"},
    {"heading": "Aldir Passarinho (1987-1989)",
     "file": "Plaqueta_Possepresidencial_AldirPassarinho_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 1987",
     "presidente": "Aldir Guimar\u00e3es Passarinho",
     "vice": ""},
    {"heading": "N\u00e9ri da Silveira (1989-1991)",
     "file": "Plaqueta_Possepresidencial_NeridaSilveira_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 14 de mar\u00e7o de 1989",
     "presidente": "Jos\u00e9 N\u00e9ri da Silveira",
     "vice": "Aldir Guimar\u00e3es Passarinho"},
    {"heading": "Oct\u00e1vio Gallotti (1991-1993)",
     "file": "Plaqueta_Possepresidencial_OctavioGallotti_NOVACAPA.txt",
     "data": "Sess\u00e3o solene de 1991",
     "presidente": "Oct\u00e1vio Gallotti",
     "vice": ""},
]

# Termos de posse manuscritos transcritos (1963-1979)
termos_json = "C:/Users/medin/Desktop/backup_judx/resultados/termos_posse_stf_1963_1979.json"

def add_plaqueta_to_doc(doc, plq):
    with open(os.path.join(ocr_dir, plq['file']), 'r', encoding='utf-8') as f:
        raw = f.read()
    cleaned = clean_ocr(raw)
    sections = extract_sections(cleaned)
    p = doc.add_paragraph()
    run = p.add_run(plq['presidente'].upper())
    run.font.size = docx.shared.Pt(14)
    meta = doc.add_paragraph()
    run = meta.add_run(plq['data'])
    run.italic = True
    run.font.size = docx.shared.Pt(11)
    for sec in sections:
        if sec['speaker']:
            p = doc.add_paragraph()
            run = p.add_run(sec['speaker'])
            run.bold = True
        paras = text_to_paragraphs(sec['text'])
        for pt in paras:
            doc.add_paragraph(pt, style='Normal')
    doc.add_paragraph('')

# Build consolidated docx: plaquetas OCR + lacunas + existing speeches
print("Criando docx consolidado completo...")
existing = docx.Document("C:/Users/medin/Downloads/discurso presidentes.docx")
consolidated = docx.Document()
style = consolidated.styles['Normal']
style.font.name = 'Calibri'
style.font.size = docx.shared.Pt(11)

# Title
p = consolidated.add_paragraph()
run = p.add_run('DISCURSOS DE POSSE NA PRESIDENCIA DO STF')
run.font.size = docx.shared.Pt(16)
run.bold = True
p = consolidated.add_paragraph()
run = p.add_run('Corpus consolidado: 1975-2025')
run.font.size = docx.shared.Pt(12)
run.italic = True
p = consolidated.add_paragraph()
run = p.add_run('Fonte: plaquetas institucionais STF (OCR via Wayback Machine) + arquivo pessoal Damares Medina')
run.font.size = docx.shared.Pt(9)
consolidated.add_paragraph('')

# Add termos de posse manuscritos (1963-1979)
import json as jsonlib
with open(termos_json, 'r', encoding='utf-8') as f:
    termos_data = jsonlib.load(f)
termos_sorted = sorted(termos_data, key=lambda t: t['data_posse'].split('/')[::-1])

print("  Adicionando 7 termos de posse manuscritos (1963-1979)...")
for t in termos_sorted:
    p = consolidated.add_paragraph()
    run = p.add_run(t['ministro'].upper())
    run.font.size = docx.shared.Pt(14)
    meta = consolidated.add_paragraph()
    run = meta.add_run(f"Posse: {t['data_posse']} | Sessao presidida por: {t['presidente_sessao']}")
    run.italic = True
    run.font.size = docx.shared.Pt(10)
    meta2 = consolidated.add_paragraph()
    run = meta2.add_run(f"Fonte: {t.get('tipo', 'Manuscrito original - livro de termos do STF')}")
    run.italic = True
    run.font.size = docx.shared.Pt(9)
    for para in t['transcricao'].split('\n\n'):
        if para.strip():
            consolidated.add_paragraph(para.strip(), style='Normal')
    if t.get('signatarios'):
        p = consolidated.add_paragraph()
        run = p.add_run('Signatarios: ')
        run.bold = True
        run.font.size = docx.shared.Pt(9)
        run2 = p.add_run(', '.join(t['signatarios']))
        run2.font.size = docx.shared.Pt(9)
    consolidated.add_paragraph('')
    print(f"    + {t['ministro']} ({t['data_posse']})")

# Add all 8 plaquetas (OCR)
print("  Adicionando 8 plaquetas OCR...")
for plq in plaquetas:
    print(f"    + {plq['heading']}")
    add_plaqueta_to_doc(consolidated, plq)

# Add existing speeches (post-Gallotti, from discurso presidentes.docx)
print("  Adicionando discursos existentes...")
for para in existing.paragraphs:
    new_para = consolidated.add_paragraph()
    try:
        new_para.style = consolidated.styles[para.style.name]
    except KeyError:
        new_para.style = consolidated.styles['Normal']
    for run in para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        if run.font.size:
            new_run.font.size = run.font.size

consolidated_path = "C:/Users/medin/Downloads/Discursos posse STF - consolidado 1975-2025.docx"
consolidated.save(consolidated_path)
print(f"\n  Consolidado: {consolidated_path}")
print(f"  Tamanho: {os.path.getsize(consolidated_path):,} bytes")
print(f"\n  8 plaquetas OCR + discursos existentes + 7 lacunas documentadas")
